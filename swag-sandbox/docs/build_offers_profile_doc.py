"""
Build the 'dtcmvp SWAG Directory' team planning document as .docx.
Run: python3 build_offers_profile_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

GREEN = RGBColor(0x2E, 0xD5, 0x73)
BLUE = RGBColor(0x1E, 0x90, 0xFF)
ORANGE = RGBColor(0xFF, 0x9F, 0x43)
DARK = RGBColor(0x2F, 0x35, 0x42)
MUTED = RGBColor(0x71, 0x80, 0x96)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Helvetica Neue'
style.font.size = Pt(11)
style.font.color.rgb = DARK
for section in doc.sections:
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)


def add_title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(2)

def add_subtitle(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED
    run.italic = True
    p.paragraph_format.space_after = Pt(18)

def add_h1(text, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = color or GREEN
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)

def add_h2(text, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = color or DARK
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def add_h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

def add_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(6)

def add_bullet(text, bold_lead=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_lead:
        r = p.add_run(bold_lead + ' ')
        r.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)

def add_callout(label, body, color):
    p = doc.add_paragraph()
    r1 = p.add_run(label + '  ')
    r1.bold = True
    r1.font.color.rgb = color
    p.add_run(body).font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(8)

def add_code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Menlo'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def add_divider():
    p = doc.add_paragraph()
    run = p.add_run('─' * 60)
    run.font.color.rgb = MUTED
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)


# ════════════════════════════════════════════════════════════
# TITLE
# ════════════════════════════════════════════════════════════
add_title('dtcmvp SWAG Directory')
add_subtitle('internal planning doc · v1.0 · team brief')

# ════════════════════════════════════════════════════════════
# WHAT IS THIS
# ════════════════════════════════════════════════════════════
add_h1('What is the SWAG Directory')
add_body(
    'dtcmvp publishes a first-party ROI analysis for Shopify apps in our directory. For each '
    'partner, we build a page that tells a brand — in their own numbers — what that app is '
    'worth to their store. We read the partner\'s public website, apply our own formulas and '
    'industry SWAGs (scientific wild-ass guesses), and render a personalized answer using the '
    'brand\'s profile data. The partner does not need to participate.'
)
add_body(
    'The key move: instead of showing the partner\'s price and asking "is this a good deal?", '
    'we ask the brand what ROI multiple they want on their tools (5x, 8x, 15x) and tell them '
    'the maximum they should pay to hit that target. Pricing is optional. Power stays with the '
    'brand.'
)

# ════════════════════════════════════════════════════════════
# HOW IT WORKS
# ════════════════════════════════════════════════════════════
add_h1('How it works (short version)')

add_h2('Three inputs to any SWAG page')
add_bullet('Brand enters once during onboarding: annual orders, AOV, return rate, avg cost per item, category.', bold_lead='Brand profile —')
add_bullet('Industry averages we ship as defaults ($2/ticket, 1.5% edit rate, etc.). Every SWAG is editable — the brand can override any number if they know theirs better.', bold_lead='dtcmvp SWAGs —')
add_bullet('What the vendor claims to do, extracted from their public website. We map each claim to a formula, apply brand numbers + SWAGs, and compute annual value.', bold_lead='Partner benefits —')

add_h2('Sourcing tiers')
add_bullet('Default. We read the partner\'s public marketing and apply our own math. ~15 min per partner.', bold_lead='Free listing:')
add_bullet('Partner has a live ROI calculator on their site (like Aftersell does) — we can absorb their published math and cite "using the partner\'s own methodology." Can be a free upgrade when we find a good one, or a paid premium feature. ~45 min.', bold_lead='Pro listing (partner-aligned):')
add_bullet('Partner gives us private data (case studies, real client reports, call transcripts for dtcmvp clients). Highest confidence. Premium or client-only.', bold_lead='Pro listing (partner-authored):')

add_h2('Commercial model')
add_callout('Brand:', '$50 per booked call with a partner, paid once the meeting happens. Free access to the SWAG Directory.', GREEN)
add_callout('Partner:', 'Free listing = dtcmvp-authored SWAG. Pro listing = partner-aligned or partner-authored math, paid tier.', BLUE)
add_callout('dtcmvp:', '10% of deals that close through this channel. SWAG Directory is the top of funnel.', ORANGE)

# ════════════════════════════════════════════════════════════
# JOBS TO BE DONE
# ════════════════════════════════════════════════════════════
add_h1('Jobs to be done', color=BLUE)

add_h2('Sean · SWAG creation backend', color=GREEN)
add_body(
    'Sean builds (with Claude Code) a backend admin workflow for turning any Shopify vendor\'s '
    'public website into a SWAG spec. This is the factory — every premium listing starts here.'
)
add_bullet('Build a Claude Code skill that can be pointed at any Shopify app website and generates a structured SWAG: vendor name, benefits identified, formulas, brand inputs needed, SWAG defaults, confidence tags.')
add_bullet('Admin page at /admin that loads all partner listings from the directory. Each partner has a status badge (not started / drafting / published). Click one to start the SWAG process.')
add_bullet('Creation happens in the terminal: Sean picks a partner, Claude fetches their site, proposes a SWAG, Sean reviews/iterates, approves. The page goes live as a static calculator.')
add_bullet('Backfill at scale — goal is to have a critical mass of premium SWAG listings before launch. This is the gating item for go-live.')
add_bullet('Co-own the marketing launch: social media pushes with dtcmvp brands + partners, potentially an ebook or co-branded content piece.')

add_h2('Andrew · Brand experience + design', color=ORANGE)
add_body(
    'Andrew owns everything the brand sees — from onboarding through the SWAG page to the '
    'booking CTA.'
)
add_bullet('Brand onboarding flow — a form where brands enter their profile once (annual orders, AOV, return rate, avg cost per item, category). This data powers every SWAG in the directory.')
add_bullet('Refactor free vs. pro listings. Not every listing gets a SWAG (some vendor types won\'t have a calculable ROI). A "pro" listing has the SWAG page; a "free" listing is a standard directory card. The visual distinction needs to be clear and aspirational — pro should feel premium.')
add_bullet('SWAG page design — the actual calculator layout. Brand profile on the left, SWAGs in the middle (editable), computed output on the right, target ROI at the top. See the POC at localhost:3009/calculators/orderediting for the current state in dtcmvp design language.')
add_bullet('Canva / video assets for the marketing launch. We need visuals that explain the concept to brands + partners in <30 seconds.')
add_bullet('"AI is analyzing your data" loading animation for the SWAG compute step. Short, branded, fun.')

add_h2('Peter · Backend integration + deal flow', color=BLUE)
add_body(
    'Peter wires up the plumbing that connects SWAGs to real business outcomes — saving profile '
    'data, processing claims, and closing the loop on meetings and payments.'
)
add_bullet('Make sure the Slack / Airtable / app process all comes together. The Airtable "Offers" table is the v1 placeholder for partner listings (may move, Sean will flag). Brand profile fields need a home — existing brand record or new schema.')
add_bullet('SWAG saving — when a brand views a SWAG page, their computed results get stored so we have a record of what they saw and when.')
add_bullet('SWAG claim flow — brand clicks the CTA to book a meeting. This triggers a confirmation check (v1: Sean manually pings the partner; v2: auto-qualification based on partner criteria).')
add_bullet('Meeting + payment flow — once the meeting is confirmed and held, brand gets $50 via Tremendous (same infra as feedback calls). dtcmvp tracks the 10% rev share on any resulting deal.')
add_bullet('Attribution — how do we know a closed deal came from the SWAG Directory? Needs a tracking param or UTM in the booking flow. This is what the rev share contract hangs on.')

# ════════════════════════════════════════════════════════════
# TIMELINE
# ════════════════════════════════════════════════════════════
add_h1('Timeline')

add_h2('Phase 0 · Foundation (now)')
add_bullet('Sean builds the SWAG skill in Claude Code and proves it on 3-5 partners end-to-end.')
add_bullet('Andrew reviews the OE calculator POC and starts designing the brand onboarding flow + free vs pro listing treatment.')
add_bullet('Peter confirms the Airtable data model and starts scoping the claim + payment flow.')

add_h2('Phase 1 · Build')
add_bullet('Sean backfills 15-20+ pro listings using the skill. This is serial work but each one is ~15-45 min.')
add_bullet('Andrew ships brand onboarding and the SWAG page design in the directory app.')
add_bullet('Peter ships SWAG saving, claim flow, meeting confirmation, and $50 payout.')

add_h2('Phase 2 · Soft launch')
add_bullet('Invite 10-20 pilot brands to create profiles and explore the SWAG Directory.')
add_bullet('Measure: how many brands create profiles? How many view SWAGs? How many book calls? How many result in deals?')
add_bullet('Iterate on SWAG quality, onboarding friction, claim flow based on pilot feedback.')

add_h2('Phase 3 · Marketing launch')
add_bullet('Andrew creates Canva/video launch assets.')
add_bullet('Sean + dtcmvp brands co-promote on social media.')
add_bullet('Partners get notified their listings have been upgraded — invitation to provide evidence for Tier 1/2 upgrades.')
add_bullet('Potentially an ebook or co-branded content piece with launch partners.')

add_h2('Later')
add_bullet('AI-curated directory: "brands like yours got the most from these tools" using profile + deal history.')
add_bullet('Partner self-serve evidence upload for pro listing upgrades.')
add_bullet('Auto-qualification: partners set matching criteria, good-fit brands get instant meeting approval.')
add_bullet('Variable bounties per partner. Dynamic pricing replaces flat $50.')

# ════════════════════════════════════════════════════════════
# APPENDIX
# ════════════════════════════════════════════════════════════
add_divider()
add_h1('Appendix · Worked example: Order Editing', color=MUTED)
add_body(
    'This is the full exercise for building a SWAG from scratch. No partner call, no sheet, '
    'no private data — just orderediting.com and dtcmvp\'s brain. This is what the Claude Code '
    'skill will automate.'
)

add_h2('Step 1 · Read the website')
add_body(
    'We pulled orderediting.com and found two quantifiable benefits:'
)
add_bullet('"98% reduction in order-change enquiries" (Oh Polly case study) + "20% reduction in total CS tickets."', bold_lead='Support reduction:')
add_bullet('Named customer results — Origin USA $22k upsell in 2 months, David Protein ~$1k/day, Nakie $80k/month.', bold_lead='Upsell revenue:')
add_body(
    'Their site does not mention return prevention or processing-fee recovery, even though '
    'their internal sheet values both. Public info finds 2 of 4 real benefits — that\'s fine. '
    'The SWAG is a floor, not a ceiling.'
)

add_h2('Step 2 · Map each benefit to inputs')
add_body(
    'For each benefit: what does a brand need to know to put a dollar on it? Which inputs '
    'come from their profile, which do we SWAG? Every SWAG is a default the brand can override.'
)

add_h3('Support ticket reduction')
add_code_block(
    'Input                           Source            Default    Editable?\n'
    '─────────────────────────────   ───────────────   ────────   ─────────\n'
    'Annual order volume             brand profile     —          yes\n'
    'Edit-ticket rate (% of orders)  dtcmvp SWAG       1.5%       yes\n'
    'Cost per support ticket         dtcmvp SWAG       $2.00      yes\n'
    'Reduction achieved              dtcmvp SWAG       80%        yes\n'
    '\n'
    'formula:  orders × 0.015 × $2.00 × 0.80\n'
    'example:  250,000 × 0.015 × 2 × 0.80 = $6,000/yr'
)

add_h3('Post-purchase upsell')
add_code_block(
    'Input                           Source            Default    Editable?\n'
    '─────────────────────────────   ───────────────   ────────   ─────────\n'
    'Annual order volume             brand profile     —          yes\n'
    'AOV                             brand profile     —          yes\n'
    'Upsell take rate                dtcmvp SWAG       1.5%       yes\n'
    'Upsell as % of AOV              dtcmvp SWAG       40%        yes\n'
    '\n'
    'formula:  orders × 0.015 × (AOV × 0.40)\n'
    'example:  250,000 × 0.015 × ($120 × 0.40) = $180,000/yr'
)

add_h2('Step 3 · Roll up + target ROI flip')
add_code_block(
    'For a 250k-order brand @ $120 AOV:\n'
    '\n'
    '  Support reduction   $6,000\n'
    '  Upsell revenue      $180,000\n'
    '  ──────────────────────────────\n'
    '  Total annual value  $186,000\n'
    '\n'
    '  Brand target: 8x ROI\n'
    '  → max monthly price: $186,000 / 8 / 12 ≈ $1,938/mo\n'
    '\n'
    '  If we know pricing ($800/mo): "That\'s 19x. Strong fit."\n'
    '  If we don\'t:                 "Ask for $1,938/mo or less."'
)

add_h2('Step 4 · Publish')
add_body(
    'The spec gets written to a partner config file. The calculator page auto-renders from it '
    'in dtcmvp design language. Sean reviews, approves, status flips to published. Done.'
)

# ─── footer ────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
run = p.add_run('dtcmvp · SWAG Directory · v1.0 · internal')
run.font.size = Pt(9)
run.font.color.rgb = MUTED
run.italic = True

out = '/Users/seanwendt/Documents/Payments/roi-carwash/docs/dtcmvp-offers-profile-plan.docx'
doc.save(out)
print('wrote', out)
